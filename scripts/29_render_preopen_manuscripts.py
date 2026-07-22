#!/usr/bin/env python3
"""Render the PRE-OPEN manuscript DOCX files from their canonical Markdown.

This helper intentionally performs no claim inference.  It formats the current
PRE-OPEN prose, removes machine-only HTML claim comments from visible Word
content, preserves the eight scope sentences as ordinary paragraphs, and fails
if a known legacy result claim survives into an artifact.

Run with the Codex primary runtime Python because it provides python-docx:

    /path/to/codex-primary-runtime/dependencies/python/bin/python \
        scripts/29_render_preopen_manuscripts.py
"""
# ruff: noqa: E402 -- the PRE guard must run before any python-docx import.

from __future__ import annotations

import argparse
from copy import deepcopy
import html
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from _preopen_manuscript_guard import (
    PreopenManuscriptGuardError,
    assert_preopen_manuscript_render_allowed,
)


def _guard_before_docx_import(argv: Sequence[str]) -> Path:
    """Resolve ``--root`` and enforce PRE state before importing python-docx."""
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument(
        "--root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    known, _unknown = parser.parse_known_args(argv)
    root = known.root.resolve()
    try:
        assert_preopen_manuscript_render_allowed(root)
    except PreopenManuscriptGuardError as exc:
        raise SystemExit(f"PRE-OPEN manuscript render refused: {exc}") from exc
    return root


_EARLY_GUARDED_ROOT: Path | None = None
if __name__ == "__main__":
    # This deliberately precedes every python-docx import below.
    _EARLY_GUARDED_ROOT = _guard_before_docx_import(sys.argv[1:])

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B78"
GOLD = "8A6814"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BLACK = "000000"


@dataclass(frozen=True)
class Preset:
    name: str
    body_alignment: WD_ALIGN_PARAGRAPH
    body_after_pt: float
    body_line: float
    h1_before_pt: float
    h1_after_pt: float
    h2_before_pt: float
    h2_after_pt: float
    h3_before_pt: float
    h3_after_pt: float
    list_marker_in: float
    list_text_in: float
    list_hanging_in: float
    list_after_pt: float
    list_line: float
    table_fill: str


PRESETS = {
    "narrative_proposal": Preset(
        name="narrative_proposal",
        body_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        body_after_pt=8,
        body_line=1.333,
        h1_before_pt=18,
        h1_after_pt=10,
        h2_before_pt=12,
        h2_after_pt=6,
        h3_before_pt=8,
        h3_after_pt=4,
        list_marker_in=0.181,
        list_text_in=0.375,
        list_hanging_in=0.194,
        list_after_pt=4,
        list_line=1.208,
        table_fill=CALLOUT,
    ),
    "standard_business_brief": Preset(
        name="standard_business_brief",
        body_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        body_after_pt=6,
        body_line=1.10,
        h1_before_pt=16,
        h1_after_pt=8,
        h2_before_pt=12,
        h2_after_pt=6,
        h3_before_pt=8,
        h3_after_pt=4,
        list_marker_in=0.25,
        list_text_in=0.50,
        list_hanging_in=0.25,
        list_after_pt=8,
        list_line=1.167,
        table_fill=LIGHT_GRAY,
    ),
    "compact_reference_guide": Preset(
        name="compact_reference_guide",
        body_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        body_after_pt=6,
        body_line=1.25,
        h1_before_pt=18,
        h1_after_pt=10,
        h2_before_pt=14,
        h2_after_pt=7,
        h3_before_pt=10,
        h3_after_pt=5,
        list_marker_in=0.187,
        list_text_in=0.375,
        list_hanging_in=0.188,
        list_after_pt=4,
        list_line=1.25,
        table_fill=LIGHT_BLUE,
    ),
}


SCOPE_SENTENCES = (
    "Route A is history-dependent and uses target-site water-temperature "
    "observations through each issue date; it does not establish ungauged prediction.",
    "Route A does not establish independent river-network separation, physical "
    "river-network routing, or hydraulic travel time.",
    "Route A is a one-shot retrospective historical-information evaluation, not "
    "an operational replay with archived as-issued predictor vintages or future NWP.",
    "Route A uses daily-mean statistical thresholds and a numerical non-inferiority "
    "margin; neither has ecological, biological, or regulatory meaning.",
    "Route A provides no physical, deployment, regulatory, or distribution-free "
    "safety guarantee, and failure to reject is not equivalence.",
    "Route-A architecture controls and predictor sensitivities are descriptive or "
    "exploratory and do not identify causal mechanisms.",
    "Route A evaluates a fixed availability-enriched 120-site cohort and is not "
    "nationally representative of all U.S. rivers or all calendar days.",
    "Route A does not establish conditional coverage, and its equal-weight "
    "three-quantile pinball summary is not CRPS.",
)


LEGACY_PATTERNS = {
    "legacy 40-site cohort": re.compile(r"\b40\s+(?:public\s+)?USGS stations\b", re.I),
    "legacy reportable N=114": re.compile(
        r"\b(?:n\s*=\s*)?114\s+(?:blind-test\s+)?stations\b", re.I
    ),
    "legacy spatial skill triplet": re.compile(r"\+0\.13\s*/\s*\+0\.14\s*/\s*\+0\.23"),
    "legacy RMSE triplet": re.compile(r"0\.554\s*/\s*1\.175\s*/\s*1\.490"),
    "legacy air2stream triplets": re.compile(r"0\.630\s*/\s*1\.289\s*/\s*1\.658"),
    "legacy coverage range": re.compile(r"89\s*[-–—]\s*97\s*%"),
    "legacy transfer distance": re.compile(r"\b358\s*km\b", re.I),
    "legacy blind-test wording": re.compile(
        r"\b(?:one-shot\s+)?2019\s*[-–—]\s*2020\s+blind\s*[- ]?test\b", re.I
    ),
    "legacy superiority claim": re.compile(r"\bsignificantly beats\b", re.I),
    "legacy bounded-degradation claim": re.compile(r"\bbounded-degradation guarantee\b", re.I),
    "legacy distribution-free claim": re.compile(
        r"\bdistribution-free calibrated uncertainty\b", re.I
    ),
}


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size_pt: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_paragraph_spacing(
    paragraph: Paragraph,
    *,
    before: float,
    after: float,
    line: float,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_paragraph_shading(paragraph: Paragraph, fill: str, border: str = BLUE) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    pbdr.append(left)
    ppr.append(pbdr)


def _set_keep(paragraph: Paragraph, *, keep_next: bool = False, keep_lines: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def _add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    _set_run_font(run, size_pt=9, color=MUTED)


def _configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _configure_styles(doc: DocumentType, preset: Preset) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(preset.body_after_pt)
    normal.paragraph_format.line_spacing = preset.body_line
    normal.paragraph_format.alignment = preset.body_alignment

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, preset.h1_before_pt, preset.h1_after_pt),
        ("Heading 2", 13, BLUE, preset.h2_before_pt, preset.h2_after_pt),
        ("Heading 3", 12, DARK_BLUE, preset.h3_before_pt, preset.h3_after_pt),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.line_spacing = 1.0
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name in ("PreOpen Kicker", "PreOpen Callout", "PreOpen Metadata"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    kicker = styles["PreOpen Kicker"]
    kicker.font.name = "Calibri"
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(GOLD)
    kicker._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    kicker._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    kicker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(18)
    kicker.paragraph_format.line_spacing = 1.0

    callout = styles["PreOpen Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.10)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15

    meta = styles["PreOpen Metadata"]
    meta.font.name = "Calibri"
    meta.font.size = Pt(10.5)
    meta.font.color.rgb = RGBColor.from_string(BLACK)
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.line_spacing = 1.10


def _add_numbering(doc: DocumentType, preset: Preset) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(n.get(qn("w:abstractNumId")))
        for n in numbering.findall(qn("w:abstractNum"))
        if n.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def make_abstract(abstract_id: int, *, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(round(preset.list_text_in * 1440)))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(round(preset.list_text_in * 1440)))
        ind.set(qn("w:hanging"), str(round(preset.list_hanging_in * 1440)))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), str(round(preset.list_after_pt * 20)))
        spacing.set(qn("w:line"), str(round(preset.list_line * 240)))
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend((tabs, ind, spacing))
        lvl.extend((start, num_fmt, lvl_text, suff, ppr))
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id: int, abstract_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    make_abstract(next_abstract, fmt="bullet", text="•")
    make_num(next_num, next_abstract)
    make_abstract(next_abstract + 1, fmt="decimal", text="%1.")
    make_num(next_num + 1, next_abstract + 1)
    return next_num, next_num + 1


def _apply_num(paragraph: Paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    ppr.append(num_pr)


def _fresh_num_instance(doc: DocumentType, template_num_id: int) -> int:
    """Create a renderer-independent list definition that restarts at one."""
    numbering = doc.part.numbering_part.element
    template = next(
        (
            node
            for node in numbering.findall(qn("w:num"))
            if node.get(qn("w:numId")) == str(template_num_id)
        ),
        None,
    )
    if template is None:
        raise AssertionError(f"numbering template not found: {template_num_id}")
    abstract_ref = template.find(qn("w:abstractNumId"))
    if abstract_ref is None or abstract_ref.get(qn("w:val")) is None:
        raise AssertionError(f"numbering template has no abstractNumId: {template_num_id}")
    template_abstract_id = str(abstract_ref.get(qn("w:val")))
    template_abstract = next(
        (
            node
            for node in numbering.findall(qn("w:abstractNum"))
            if node.get(qn("w:abstractNumId")) == template_abstract_id
        ),
        None,
    )
    if template_abstract is None:
        raise AssertionError(
            f"numbering template has no abstract definition: {template_num_id}"
        )
    used_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(used_abstract, default=0) + 1
    abstract = deepcopy(template_abstract)
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = abstract.find(qn("w:nsid"))
    if nsid is not None:
        nsid.set(qn("w:val"), f"F{abstract_id:07X}"[-8:])
    level = abstract.find(qn("w:lvl"))
    if level is None:
        raise AssertionError("numbering abstract definition has no level zero")
    start_definition = level.find(qn("w:start"))
    if start_definition is None:
        start_definition = OxmlElement("w:start")
        level.insert(0, start_definition)
    start_definition.set(qn("w:val"), "1")
    first_num_index = next(
        (
            index
            for index, node in enumerate(numbering)
            if node.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)
    used = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(used, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def _set_cell_margins(
    cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Table, widths_dxa: Sequence[int], *, indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must sum to 9360 DXA: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_borders(table: Table, color: str = "C7CED8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def _clean_markdown_text(text: str) -> str:
    text = html.unescape(text)
    text = text.replace("\\*", "*")
    text = re.sub(r"\^([^\^]+)\^", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*))")


def _add_inline(paragraph: Paragraph, text: str, *, default_bold: bool = False) -> None:
    text = _clean_markdown_text(text)
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _set_run_font(run, bold=default_bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name="Consolas", size_pt=9.5, color=DARK_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, italic=True, bold=default_bold)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, bold=default_bold)


def _strip_machine_comments(markdown: str) -> str:
    return re.sub(r"<!--.*?-->", "", markdown, flags=re.S)


def _blocks(markdown: str) -> list[tuple[str, object]]:
    """Parse the small Markdown subset used by the three PRE-OPEN sources."""
    text = _strip_machine_comments(markdown)
    lines = text.splitlines()
    blocks: list[tuple[str, object]] = []
    i = 0
    paragraph: list[str] = []

    def flush() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(("paragraph", " ".join(x.strip() for x in paragraph)))
            paragraph = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush()
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
                rows.pop(1)
            blocks.append(("table", rows))
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            flush()
            blocks.append((f"heading{len(heading.group(1))}", heading.group(2)))
            i += 1
            continue
        if stripped.startswith(">"):
            flush()
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                quote_lines.append(lines[i].lstrip()[1:].strip())
                i += 1
            blocks.append(("quote", " ".join(quote_lines)))
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            flush()
            i += 1
            item = [bullet.group(1)]
            while i < len(lines) and lines[i].strip() and lines[i][:1].isspace():
                item.append(lines[i].strip())
                i += 1
            blocks.append(("bullet", " ".join(item)))
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            flush()
            i += 1
            item = [numbered.group(1)]
            while i < len(lines) and lines[i].strip() and lines[i][:1].isspace():
                item.append(lines[i].strip())
                i += 1
            blocks.append(("number", " ".join(item)))
            continue
        paragraph.append(stripped.removesuffix("  "))
        i += 1
    flush()
    return blocks


def _add_table(doc: DocumentType, rows: Sequence[Sequence[str]], preset: Preset) -> None:
    if not rows:
        return
    ncols = len(rows[0])
    if any(len(row) != ncols for row in rows):
        raise ValueError("ragged Markdown table")
    table = doc.add_table(rows=len(rows), cols=ncols)
    if ncols == 4:
        widths = (1944, 2016, 1728, 3672)
    elif ncols == 2:
        widths = (2700, 6660)
    else:
        base = 9360 // ncols
        widths = tuple([base] * (ncols - 1) + [9360 - base * (ncols - 1)])
    _set_table_geometry(table, widths)
    _set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                _set_cell_shading(cell, preset.table_fill)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, ncols - 1) else WD_ALIGN_PARAGRAPH.CENTER
            )
            _set_paragraph_spacing(p, before=0, after=0, line=1.10)
            _add_inline(p, value, default_bold=r_idx == 0)
            for run in p.runs:
                if r_idx == 0:
                    _set_run_font(run, size_pt=9.5, color=INK, bold=True)
                else:
                    _set_run_font(run, size_pt=9.5)
    after = doc.add_paragraph()
    _set_paragraph_spacing(after, before=0, after=4, line=1.0)


def _render_blocks(
    doc: DocumentType,
    blocks: Iterable[tuple[str, object]],
    *,
    preset: Preset,
    bullet_num_id: int,
    decimal_num_id: int,
    skip_first_heading: bool,
) -> None:
    first_heading_skipped = not skip_first_heading
    previous_kind: str | None = None
    previous_paragraph: Paragraph | None = None
    active_bullet_id = bullet_num_id
    active_decimal_id = decimal_num_id
    for kind, payload in blocks:
        if kind == "heading1" and not first_heading_skipped:
            first_heading_skipped = True
            continue
        if kind.startswith("heading"):
            level = int(kind[-1])
            # Markdown H2 is the manuscript's top body level after the cover title.
            word_level = max(1, level - 1) if skip_first_heading else level
            p = doc.add_paragraph(style=f"Heading {word_level}")
            _add_inline(p, str(payload))
            _set_keep(p, keep_next=True)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            _add_inline(p, str(payload))
            _set_keep(p)
        elif kind == "quote":
            p = doc.add_paragraph(style="PreOpen Callout")
            _add_inline(p, str(payload))
            _set_paragraph_shading(p, CALLOUT)
            _set_keep(p)
        elif kind in ("bullet", "number"):
            if kind == "bullet" and previous_kind != "bullet":
                active_bullet_id = _fresh_num_instance(doc, bullet_num_id)
            if kind == "number" and previous_kind != "number":
                active_decimal_id = _fresh_num_instance(doc, decimal_num_id)
            p = doc.add_paragraph()
            _apply_num(p, active_bullet_id if kind == "bullet" else active_decimal_id)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_spacing(
                p,
                before=0,
                after=preset.list_after_pt,
                line=preset.list_line,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            _add_inline(p, str(payload))
            _set_keep(p)
        elif kind == "table":
            if previous_kind == "paragraph" and previous_paragraph is not None:
                _set_keep(previous_paragraph, keep_next=True)
            _add_table(doc, payload, preset)  # type: ignore[arg-type]
        else:
            raise AssertionError(f"unknown block kind: {kind}")
        previous_kind = kind
        previous_paragraph = p if kind != "table" else None


def _set_section_columns(section, count: int, *, space_dxa: int = 360) -> None:
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_dxa))


def _set_running_furniture(doc: DocumentType, *, left: str, right: str) -> None:
    for section in doc.sections:
        _configure_section(section)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, before=0, after=0, line=1.0)
        left_run = p.add_run(left)
        _set_run_font(left_run, size_pt=8.5, color=MUTED, bold=True)
        tab = p.add_run("\t")
        _set_run_font(tab, size_pt=8.5)
        tabs = p._p.get_or_add_pPr().get_or_add_tabs()
        stop = OxmlElement("w:tab")
        stop.set(qn("w:val"), "right")
        stop.set(qn("w:pos"), "9360")
        tabs.append(stop)
        right_run = p.add_run(right)
        _set_run_font(right_run, size_pt=8.5, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(fp, before=0, after=0, line=1.0)
        run = fp.add_run("PRE-OPEN / INCOMPLETE  •  Page ")
        _set_run_font(run, size_pt=9, color=MUTED)
        _add_page_field(fp)


def _add_editorial_cover(
    doc: DocumentType,
    title: str,
    status_text: str,
    author_block: Sequence[str],
) -> None:
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=42, after=0, line=1.0)
    kicker = doc.add_paragraph(style="PreOpen Kicker")
    _add_inline(kicker, "PRE-OPEN RESEARCH MANUSCRIPT")
    p = doc.add_paragraph(style="Title")
    _add_inline(p, title)
    subtitle = doc.add_paragraph(style="Subtitle")
    _add_inline(subtitle, "Frozen design, audit boundary, and implementation status")
    for index, line in enumerate((*author_block, "22 July 2026")):
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(meta, before=18 if index == 0 else 0, after=4, line=1.0)
        run = meta.add_run(line)
        _set_run_font(
            run,
            size_pt=10.5 if index == 0 else 9.25,
            color=BLACK if index == 0 else MUTED,
            italic=index in (1, 2),
        )
    status = doc.add_paragraph(style="PreOpen Callout")
    _add_inline(status, status_text)
    _set_paragraph_shading(status, CALLOUT, GOLD)
    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(scope, before=10, after=0, line=1.15)
    run = scope.add_run("No current empirical performance conclusion is claimed.")
    _set_run_font(run, size_pt=10.5, color=INK, bold=True)
    doc.add_page_break()


def _add_memo_masthead(
    doc: DocumentType,
    *,
    title: str,
    subtitle: str,
    metadata: Sequence[tuple[str, str]],
) -> None:
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, before=14, after=0, line=1.0)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        _set_run_font(run, size_pt=23, color=BLACK, bold=True)
    _add_inline(p, title)
    for run in p.runs:
        _set_run_font(run, size_pt=23, color=BLACK, bold=True)
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.paragraph_format.space_after = Pt(14)
    _add_inline(sub, subtitle)
    for label, value in metadata:
        mp = doc.add_paragraph(style="PreOpen Metadata")
        label_run = mp.add_run(f"{label}: ")
        _set_run_font(label_run, size_pt=10.5, color=BLACK, bold=True)
        value_run = mp.add_run(value)
        _set_run_font(value_run, size_pt=10.5, color=BLACK)
    rule = doc.add_paragraph()
    _set_paragraph_spacing(rule, before=7, after=10, line=1.0)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), DARK_BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _new_document(preset: Preset) -> tuple[DocumentType, int, int]:
    doc = Document()
    # Remove the template's empty body paragraph; builders add deliberate content.
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    for section in doc.sections:
        _configure_section(section)
    _configure_styles(doc, preset)
    bullet_num, decimal_num = _add_numbering(doc, preset)
    return doc, bullet_num, decimal_num


def _build_main(root: Path) -> Path:
    source = root / "paper/ThermoRoute_paper.md"
    output = root / "paper/ThermoRoute_paper.docx"
    markdown = source.read_text(encoding="utf-8")
    blocks = _blocks(markdown)
    title_block = next(value for kind, value in blocks if kind == "heading1")
    status = next(value for kind, value in blocks if kind == "quote")
    title_idx = next(i for i, (kind, _) in enumerate(blocks) if kind == "heading1")
    status_idx = next(i for i, (kind, _) in enumerate(blocks) if kind == "quote")
    author_block = [
        _clean_markdown_text(str(value))
        for kind, value in blocks[title_idx + 1 : status_idx]
        if kind == "paragraph"
    ]
    # Start generic body at Abstract; author/status material is represented on cover.
    abstract_idx = next(
        i for i, (kind, value) in enumerate(blocks) if kind == "heading2" and value == "Abstract"
    )
    body_blocks = blocks[abstract_idx:]

    preset = PRESETS["narrative_proposal"]
    doc, bullet_num, decimal_num = _new_document(preset)
    doc.core_properties.title = _clean_markdown_text(str(title_block))
    doc.core_properties.subject = "PRE-OPEN / INCOMPLETE manuscript"
    doc.core_properties.comments = (
        "Generated from canonical PRE-OPEN Markdown; no current performance result."
    )
    _set_running_furniture(doc, left="THERMOROUTE", right="PRE-OPEN RESEARCH MANUSCRIPT")
    _add_editorial_cover(
        doc,
        _clean_markdown_text(str(title_block)),
        _clean_markdown_text(str(status)),
        author_block,
    )
    _render_blocks(
        doc,
        body_blocks,
        preset=preset,
        bullet_num_id=bullet_num,
        decimal_num_id=decimal_num,
        # The title was moved to the editorial cover; Markdown H2 becomes Word H1.
        skip_first_heading=True,
    )
    doc.save(output)
    return output


def _build_cover_letter(root: Path) -> Path:
    source = root / "paper/cover_letter.md"
    output = root / "paper/cover_letter.docx"
    blocks = _blocks(source.read_text(encoding="utf-8"))
    preset = PRESETS["standard_business_brief"]
    doc, bullet_num, decimal_num = _new_document(preset)
    doc.core_properties.title = "ThermoRoute draft cover letter"
    doc.core_properties.subject = "PRE-OPEN / INCOMPLETE - not ready for submission"
    doc.core_properties.comments = "Generated from canonical PRE-OPEN Markdown."
    _set_running_furniture(doc, left="THERMOROUTE | COVER LETTER", right="NOT FOR SUBMISSION")
    _add_memo_masthead(
        doc,
        title="DRAFT COVER LETTER",
        subtitle="NOT READY FOR SUBMISSION - PRE-OPEN / INCOMPLETE",
        metadata=(
            ("Status", "Internal preparation draft; do not send to a journal"),
            ("Manuscript", "ThermoRoute retrospective daily river-temperature hindcasting"),
            ("Date", "22 July 2026"),
        ),
    )
    _render_blocks(
        doc,
        blocks,
        preset=preset,
        bullet_num_id=bullet_num,
        decimal_num_id=decimal_num,
        skip_first_heading=True,
    )
    doc.save(output)
    return output


def _build_highlights(root: Path) -> Path:
    source = root / "paper/highlights.md"
    output = root / "paper/highlights.docx"
    blocks = _blocks(source.read_text(encoding="utf-8"))
    preset = PRESETS["compact_reference_guide"]
    doc, bullet_num, decimal_num = _new_document(preset)
    doc.core_properties.title = "ThermoRoute PRE-OPEN highlights"
    doc.core_properties.subject = "PRE-OPEN / INCOMPLETE design highlights"
    doc.core_properties.comments = (
        "Generated from canonical PRE-OPEN Markdown; design highlights only."
    )
    _set_running_furniture(doc, left="THERMOROUTE | HIGHLIGHTS", right="DESIGN HIGHLIGHTS ONLY")
    _add_memo_masthead(
        doc,
        title="THERMOROUTE HIGHLIGHTS",
        subtitle="PRE-OPEN / INCOMPLETE - DESIGN HIGHLIGHTS ONLY",
        metadata=(
            ("Evidence state", "Canonical computation and target-period evaluation incomplete"),
            ("Use", "Internal review; not a performance summary"),
        ),
    )
    content = blocks[1:] if blocks and blocks[0][0] == "heading1" else blocks
    first_bullet = next(i for i, (kind, _) in enumerate(content) if kind == "bullet")
    after_bullets = first_bullet
    while after_bullets < len(content) and content[after_bullets][0] == "bullet":
        after_bullets += 1
    _render_blocks(
        doc,
        content[:first_bullet],
        preset=preset,
        bullet_num_id=bullet_num,
        decimal_num_id=decimal_num,
        skip_first_heading=True,
    )
    two_col = doc.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(two_col)
    _set_section_columns(two_col, 2)
    _render_blocks(
        doc,
        content[first_bullet:after_bullets],
        preset=preset,
        bullet_num_id=bullet_num,
        decimal_num_id=decimal_num,
        skip_first_heading=True,
    )
    one_col = doc.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(one_col)
    _set_section_columns(one_col, 1)
    _render_blocks(
        doc,
        content[after_bullets:],
        preset=preset,
        bullet_num_id=bullet_num,
        decimal_num_id=decimal_num,
        skip_first_heading=True,
    )
    doc.save(output)
    return output


def _docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(parts)


def _audit_docx(
    path: Path,
    *,
    preset: Preset,
    require_scope_sentences: bool,
) -> None:
    text = _docx_text(path)
    folded = text.casefold()
    if "pre-open" not in folded or "incomplete" not in folded:
        raise AssertionError(f"{path}: missing PRE-OPEN / incomplete disclosure")
    for name, pattern in LEGACY_PATTERNS.items():
        if pattern.search(text):
            raise AssertionError(f"{path}: prohibited {name}")
    if require_scope_sentences:
        for sentence in SCOPE_SENTENCES:
            if sentence not in text:
                raise AssertionError(f"{path}: scope sentence missing: {sentence}")

    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        if "word/comments.xml" in names:
            raise AssertionError(f"{path}: comments.xml must not exist")
        xml_text = "\n".join(
            zf.read(name).decode("utf-8", errors="replace")
            for name in names
            if name.endswith(".xml")
        )
        for marker in ("ROUTE_A_CLAIM_ENTRY", "ROUTE_A_CLAIM ", "<!--", "-->"):
            if marker in xml_text:
                raise AssertionError(f"{path}: machine-only claim metadata survived: {marker}")

        document_xml = zf.read("word/document.xml").decode("utf-8")
        for expected in (
            'w:w="12240"',
            'w:h="15840"',
            'w:top="1440"',
            'w:right="1440"',
            'w:bottom="1440"',
            'w:left="1440"',
        ):
            if expected not in document_xml:
                raise AssertionError(f"{path}: missing page geometry token {expected}")

    doc = Document(path)
    normal = doc.styles["Normal"]
    if normal.font.name != "Calibri" or normal.font.size != Pt(11):
        raise AssertionError(f"{path}: Normal font does not match {preset.name}")
    if normal.paragraph_format.space_after != Pt(preset.body_after_pt):
        raise AssertionError(f"{path}: Normal spacing does not match {preset.name}")
    line_spacing = normal.paragraph_format.line_spacing
    if not isinstance(line_spacing, float) or abs(line_spacing - preset.body_line) > 1 / 240:
        raise AssertionError(f"{path}: Normal line spacing does not match {preset.name}")
    if normal.paragraph_format.alignment != preset.body_alignment:
        raise AssertionError(f"{path}: Normal alignment does not match {preset.name}")
    for section in doc.sections:
        expected_geometry = (
            Inches(8.5),
            Inches(11),
            Inches(1),
            Inches(1),
            Inches(1),
            Inches(1),
            Inches(0.492),
            Inches(0.492),
        )
        actual_geometry = (
            section.page_width,
            section.page_height,
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
            section.header_distance,
            section.footer_distance,
        )
        if any(
            actual is None or abs(int(actual) - int(expected)) > 635
            for actual, expected in zip(actual_geometry, expected_geometry, strict=True)
        ):
            raise AssertionError(f"{path}: section geometry does not match preset baseline")
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        if (
            tbl_w is None
            or tbl_w.get(qn("w:w")) != "9360"
            or tbl_w.get(qn("w:type")) != "dxa"
            or tbl_ind is None
            or tbl_ind.get(qn("w:w")) != "120"
            or layout is None
            or layout.get(qn("w:type")) != "fixed"
        ):
            raise AssertionError(f"{path}: table geometry is not fixed 9360-DXA preset geometry")


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="repository root",
    )
    parser.add_argument(
        "--check", action="store_true", help="audit existing DOCX files without rebuilding"
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    root = args.root.resolve()
    if _EARLY_GUARDED_ROOT != root:
        assert_preopen_manuscript_render_allowed(root)
    outputs: tuple[tuple[Path, Preset], ...] = (
        (root / "paper/ThermoRoute_paper.docx", PRESETS["narrative_proposal"]),
        (root / "paper/cover_letter.docx", PRESETS["standard_business_brief"]),
        (root / "paper/highlights.docx", PRESETS["compact_reference_guide"]),
    )
    if not args.check:
        outputs = (
            (_build_main(root), PRESETS["narrative_proposal"]),
            (_build_cover_letter(root), PRESETS["standard_business_brief"]),
            (_build_highlights(root), PRESETS["compact_reference_guide"]),
        )
    for path, preset in outputs:
        _audit_docx(
            path,
            preset=preset,
            require_scope_sentences=path.name == "ThermoRoute_paper.docx",
        )
        print(f"PASS {path.relative_to(root)} ({path.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
